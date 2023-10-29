"""
This is a tool to change or adjust format 
in ftext made by bintext.py
    v0.2, developed by devseed
"""

import os
import re
import codecs
import argparse
import json
from io import StringIO
from csv import DictWriter, DictReader
from docx import Document # pip install python-docx
from docx.shared import Pt
from typing import Union, List, Dict

FTEXT_VERSION = 200

# util functions
def dump_ftext(ftexts1:List[Dict[str,Union[int,str]]], 
    ftexts2: List[Dict[str, Union[int, str]]], 
    outpath: str="", *, num_width=5, 
    addr_width=6, size_width=3) -> List[str]:
    """
    ftexts1, ftexts2 -> ftext lines
    text dict is as {'addr':, 'size':, 'text':}
    :param ftexts1[]: text dict array in '○' line, 
    :param ftexts2[]: text dict array in '●' line
    :return: ftext lines
    """

    if num_width==0:
        num_width = len(str(len(ftexts1)))
    if addr_width==0:
        d = max([t['addr'] for t in ftexts1])
        addr_width = len(hex(d))-2
    if size_width==0:
        d = max([t['size'] for t in ftexts1])
        size_width = len(hex(d))-2

    fstr1 = "○{num:0"+ str(num_width) + "d}|{addr:0" + str(addr_width) + "X}|{size:0"+ str(size_width) + "X}○ {text}\n"
    fstr2 = fstr1.replace('○', '●')
    lines = []

    length = 0
    if ftexts1 == None: 
        length = len(ftexts2)
        fstr2 += '\n'
    if ftexts2 == None: 
        length = len(ftexts1)
        fstr1 += '\n'
    if ftexts1 != None and ftexts2 != None : 
        length = min(len(ftexts1), len(ftexts2))
        fstr2 += '\n'

    for i in range(length):
        if ftexts1 != None:
            t1 = ftexts1[i]
            lines.append(fstr1.format(
                num=i,addr=t1['addr'],size=t1['size'],text=t1['text']))
        if ftexts2 != None:
            t2 = ftexts2[i]
            lines.append(fstr2.format(
                num=i,addr=t2['addr'],size=t2['size'],text=t2['text']))

    if outpath != "":
        with codecs.open(outpath, 'w', 'utf-8') as fp:
            fp.writelines(lines)
    return lines 

def load_ftext(ftextobj: Union[str, List[str]], 
    only_text = False ) -> List[Dict[str, Union[int, str]]]:
    """
    ftext lines  -> ftexts1, ftexts2
    text dict is as {'addr':, 'size':, 'text':}
    :param inobj: can be path, or lines[] 
    :return: ftexts1[]: text dict array in '○' line, 
             ftexts2[]: text dict array in '●' line
    """

    ftexts1, ftexts2 = [], []
    if type(ftextobj) == str: 
        with codecs.open(ftextobj, 'r', 'utf-8') as fp: 
            lines = fp.readlines()
    else: lines = ftextobj

    if only_text == True: # This is used for merge_text
        re_line1 = re.compile(r"^○(.+?)○[ ](.*)")
        re_line2 = re.compile(r"^●(.+?)●[ ](.*)")
        for line in lines:
            line = line.strip("\n").strip('\r')
            m = re_line1.match(line)
            if m is not None:
                ftexts1.append({'addr':0,'size':0,'text': m.group(2)})
            m = re_line2.match(line)
            if m is not None:
                ftexts2.append({'addr':0,'size':0,'text': m.group(2)})
    else:
        re_line1 = re.compile(r"^○(\d*)\|(.+?)\|(.+?)○[ ](.*)")
        re_line2 = re.compile(r"^●(\d*)\|(.+?)\|(.+?)●[ ](.*)")
        for line in lines:
            line = line.strip("\n").strip('\r')
            m = re_line1.match(line)
            if m is not None:
                ftexts1.append({'addr':int(m.group(2),16),
                'size':int(m.group(3),16),'text': m.group(4)})
            m = re_line2.match(line)
            if m is not None:
                ftexts2.append({'addr':int(m.group(2),16),
                'size':int(m.group(3),16),'text': m.group(4)})
    return ftexts1, ftexts2

def file2lines(func, argidx=0, encoding="utf-8"):
    def wrapper(*args, **kw):
        if type(args[argidx]) == str: 
            args = list(args)
            path = args[argidx]
            if path!="":
                with codecs.open(path, 'r', encoding) as fp: 
                    args[argidx] = fp.readlines()
            else: args[argidx] = None
        return func(*args, **kw)
    return wrapper

def file2bytes(func, argidx=0):
    def wrapper(*args, **kw):
        if type(args[argidx]) == str: 
            with open(args[argidx], 'rb') as fp: 
                args[argidx] = fp.read()
        return func(*args, **kw)
    return wrapper

# ftextcvt functions
def docx2ftext(wordpath, outpath="") -> List[str]:
    
    lines = []
    document = Document(wordpath)
    for p in document.paragraphs:
        # text is the whole text, p.run are every str in styles
        line = p.text.rstrip('\n').rstrip('\r') 
        lines.append(line + '\n')
    if outpath!="": 
        with open(outpath, "wb") as fp:
            for line in lines:
                fp.write(line.encode('utf-8'))
    return lines

@file2lines
def csv2ftext(csvobj: Union[str, List[str]], 
    outpath="") -> List[str]:
    
    ftexts1, ftexts2 = [], []
    for t in  DictReader(StringIO("".join(csvobj))):
        keyarr = t['key'].split('|')
        addr = int(keyarr[1], 16)
        size = int(keyarr[2], 16)
        ftexts1.append({'addr': addr, 
            'size': size, 'text': t['origin']})
        ftexts2.append({'addr': addr, 
            'size': size, 'text': t['translation']})
    return dump_ftext(ftexts1, ftexts2, outpath)

@file2lines
def json2ftext(jsonobj: Union[str, List[str]], 
    outpath="") -> List[str]:
    """
    This is for paratranz.cn format, 
    :jsonobj: [{key: idx|addr|size, origin: text, translation: text}]
    """

    ftexts1, ftexts2 = [], []
    for t in json.loads("".join(jsonobj)):
        keyarr = t['key'].split('|')
        addr = int(keyarr[1], 16)
        size = int(keyarr[2], 16)
        ftexts1.append({'addr': addr, 
            'size': size, 'text': t['origin']})
        ftexts2.append({'addr': addr, 
            'size': size, 'text': t['translation']})
    return dump_ftext(ftexts1, ftexts2, outpath)
    
@file2lines
def ftext2docx(ftextobj: Union[str, List[str]], 
    outpath="") -> Document:
    """
    if this function compiled by nuitka, 
    it needs default.docx file in ./docx/templates
    """
    lines = ftextobj
    document = Document()
    for i, line in enumerate(lines):
        line = line.rstrip('\n').rstrip('\r')
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)
    if outpath!="":
        document.save(outpath)
    return document

@file2lines
def ftext2csv(ftextobj: Union[str, List[str]],      
    outpath="", *, num_width=5, 
    addr_width=6, size_width=3) -> List[str]:
    
    ftexts1, ftexts2 = load_ftext(ftextobj, False)
    fstr = "{num:0"+ str(num_width) + "d}|{addr:0" + str(addr_width) + "X}|{size:0"+ str(size_width) + "X}"
    csvstrio = StringIO()
    csvwr = DictWriter(csvstrio, ['key', 'origin', 'translation'])
    csvwr.writeheader()
    for i, (t1, t2) in enumerate(zip(ftexts1, ftexts2)):
        csvwr.writerow({
            "key": fstr.format(
            num=i, addr=t2['addr'], size=t2['size']),
            "origin": t1['text'], "translation": t2['text']
        })

    csvlines = csvstrio.getvalue().splitlines(True)
    if outpath:
        with codecs.open(outpath, 'w', 'utf8') as fp:
            fp.writelines(csvlines)
    return csvlines

@file2lines
def ftext2json(ftextobj: Union[str, List[str]], 
    outpath="",*, num_width=5, 
    addr_width=6, size_width=3) -> List[str]:

    ftexts1, ftexts2 = load_ftext(ftextobj, False)
    jsonarray: List[Dict[str: str]] = []
    fstr = "{num:0"+ str(num_width) + "d}|{addr:0" + str(addr_width) + "X}|{size:0"+ str(size_width) + "X}"
    for i, (t1, t2) in enumerate(zip(ftexts1, ftexts2)):
        jsonarray.append({
            "key": fstr.format(
                num=i, addr=t2['addr'], size=t2['size']),
            "origin": t1['text'], "translation": t2['text']
        })
    
    jsonstr = json.dumps(jsonarray, 
        ensure_ascii=False, indent=2)
    jsonlines = jsonstr.splitlines(True)
    if outpath!="":
        with codecs.open(outpath, 'w', 'utf8') as fp:
            fp.writelines(jsonlines)
    return jsonlines

@file2lines
def formatftext(ftextobj: Union[str, List[str]], 
    outpath="") -> List[str]:

    lines = ftextobj
    for i, line in enumerate(lines):
        _c = line[0]
        if _c=='○' or _c=='●':
            _idx = line.find(_c, 1) + 1
            if _idx==-1:
                raise ValueError(f"lines[{i}] error not closed, {line}!")
            elif line[_idx]!=' ': 
                print(f"detect line[{i}] without space")
                line = line[0:_idx] + ' ' + line[_idx: ]
        line = line.rstrip('\n').rstrip('\r')
        lines[i] = line + '\n'

    if outpath!="": 
        with open(outpath, "wb") as fp:
            for line in lines:
                fp.write(line.encode('utf-8'))
    return lines

def debug():
    pass

def main(cmdstr=None):
    parser = argparse.ArgumentParser(
        description="convert the ftext format to others\n"
        "ftextcvt v0.2, developed by devseed")
    parser.add_argument("inpath", type=str)
    parser.add_argument("-o", "--outpath", 
        type=str, default="out.txt")

    if cmdstr is None: args = parser.parse_args()
    else: args = parser.parse_args(cmdstr.split(' '))
    
    outpath = args.outpath
    outpath_ext=os.path.splitext(outpath)[1].lower()
    inpath = args.inpath
    inpath_ext = os.path.splitext(inpath)[1].lower()

    if inpath_ext== '.txt': #  ftext to others
        if outpath_ext == '.txt':
            formatftext(inpath, outpath)
        elif outpath_ext == '.docx':
            ftext2docx(inpath, outpath)
        elif outpath_ext == '.csv':
            ftext2csv(inpath, outpath)
        elif outpath_ext == ".json":
            ftext2json(inpath, outpath)
        else:
            raise NotImplementedError(
                f"unkonw format {inpath_ext}->{outpath_ext}")
    elif outpath_ext == '.txt': # others to ftext
        if inpath_ext == '.docx':
            docx2ftext(inpath, outpath)
        elif inpath_ext == '.csv':
            csv2ftext(inpath, outpath)
        elif inpath_ext == '.json':
            json2ftext(inpath, outpath)
        else:
            raise NotImplementedError(
                f"unkonw format {inpath_ext}->{outpath_ext}")
        return
    else:
        raise NotImplementedError(
            f"unkonw format {inpath_ext}->{outpath_ext}")

if __name__ == '__main__':
    # debug()
    main()
    pass

"""
history:
v0.1, initial version with formatftext, docx2ftext, ftext2docx
v0.2, add support for csv and json, compatiable with paratranz.cn
"""