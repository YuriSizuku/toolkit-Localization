# -*- coding: utf-8 -*-
__description__ = """
A convert tool to change or adjust ftext
    v0.3.2, developed by devseed
"""

import os
import codecs
import argparse
import json
from io import StringIO
from glob import glob
from csv import DictWriter, DictReader
from typing import Union, List, Dict

try:
    from docx import Document # pip install python-docx
    from docx.shared import Pt
except Exception:
    Document = None
    pass

try:
    from libutil import writelines, writebytes, filter_loadfiles, ftext_t, load_ftext, save_ftext
except ImportError:
    exec("from libutil_v600 import writelines, writebytes, filter_loadfiles, ftext_t, load_ftext, save_ftext")

__version__ = 320

@filter_loadfiles((0, 'utf-8'))
def ftext2pretty(linesobj: Union[str, List[str]], outpath=None) -> List[str]:
    """
    make ftext2 format pretty, and try to fix some problem
    """

    lines = linesobj
    if len(lines) > 0: lines[0] = lines[0].lstrip("\ufeff") # remove bom
    for i, line in enumerate(lines): # try fix break
        indicator = line[0]
        line = line.rstrip('\n').rstrip('\r')
        if indicator=="○" or indicator=="●":
            text_offset = line.find(indicator, 1) + 1
            if text_offset < 0: raise ValueError(f"indicator {indicator} not closed, [lineno={i+1} line='{line}']")
            elif line[text_offset]!=" ": 
                print(f"detect no ' ' [lineno={i} line='{line}']")
                line = line[0: text_offset] + ' ' + line[text_offset:]
        lines[i] = line
    return save_ftext(*load_ftext(lines), outpath)

def ftext2csv(linesobj: Union[str, List[str]], outpath=None) -> List[str]:
    """
    convert ftext2 files to csv format
    tag(org ○, now ●),addr,size,text

    """
    
    ftexts1, ftexts2 = load_ftext(linesobj)
    assert(len(ftexts1) == len(ftexts2))
    sbufio = StringIO()
    wr = DictWriter(sbufio, ["tag", "addr", "size", "text"])
    wr.writeheader() # will automaticaly detect comma, use excel to write csv can not encode as utf8
    for i, (t1, t2) in enumerate(zip(ftexts1, ftexts2)): 
        wr.writerow({"tag": "org", "addr": hex(t1.addr), "size": hex(t1.size), "text": t1.text})
        wr.writerow({"tag": "now", "addr": hex(t2.addr), "size": hex(t2.size), "text": t2.text})
    lines = sbufio.getvalue().splitlines(True)
    sbufio.close()
    if outpath: writebytes(outpath, codecs.BOM_UTF8 + writelines(lines, "utf-8"))
    return lines

def ftext2json(ftextobj: Union[str, List[str]], outpath=None) -> List[str]:
    """
    convert ftext2 to json format
    { 
        "org": {"addr": 0, "size": 0, "text": ""}
        "now" : {"addr": 0, "size": 0, "text": ""}
    }
    """
    
    ftexts1, ftexts2 = load_ftext(ftextobj)
    assert(len(ftexts1) == len(ftexts2))
    jarr: List[Dict[str: str]] = []
    for i, (t1, t2) in enumerate(zip(ftexts1, ftexts2)):
        jarr.append({
            "org": {"addr": hex(t1.addr), "size": hex(t1.size), "text": t1.text}, 
            "now" : {"addr": hex(t2.addr), "size": hex(t2.size), "text": t2.text}
        })
    
    jstr = json.dumps(jarr, ensure_ascii=False, indent=2)
    if outpath: writebytes(outpath, jstr.encode("utf-8"))
    return jstr.splitlines(True)

def ftext2paratranz(ftextobj: Union[str, List[str]], outpath=None, width_index = (5, 6, 3)) -> List[str]:
    """
    convert ftext2 to paratranz json format
    [
        {
            "key": "num|addr|size",
            "original": "source text 原文 2",
            "translation": "translation text 译文 2"
        }
    ]
    """
    
    ftexts1, ftexts2 = load_ftext(ftextobj)
    assert(len(ftexts1) == len(ftexts2))
    
    width_num, width_addr, width_size = width_index
    keypattern = "{num:0%dd}|{addr:0%dX}|{size:0%dX}" % (width_num, width_addr, width_size)
    jarr: List[Dict[str: str]] = []
    for i, (t1, t2) in enumerate(zip(ftexts1, ftexts2)):
        jarr.append({
            "key":  keypattern.format(num=i, addr=t1.addr, size=t1.size), 
            "original" : t1.text, 
            "translation" : t2.text if t2.text != t1.text else ""
        })
    
    jstr = json.dumps(jarr, ensure_ascii=False, indent=2)
    if outpath: writebytes(outpath, jstr.encode("utf-8"))
    return jstr.splitlines(True)

@filter_loadfiles([(0, "utf-8")])
def ftext2docx(linesobj: Union[str, List[str]], outpath=None) -> Document:
    """
    if this function compiled by nuitka, 
    it needs default.docx file in ./docx/templates
    """

    lines = linesobj
    document = Document()
    for i, line in enumerate(lines):
        line = line.rstrip('\n').rstrip('\r')
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)
    if outpath: document.save(outpath)
    return document

@filter_loadfiles([(0, "utf-8")])
def csv2ftext(linesobj: Union[str, List[str]], outpath=None) -> List[str]:
    lines = linesobj
    if len(lines) > 0: lines[0] = lines[0].lstrip("\ufeff")
    ftexts1, ftexts2 = [], []
    for i, t in enumerate(DictReader(lines)):
        ftext = None
        try:
            ftext = ftext_t(int(t['addr'], 0), int(t['size'], 0), t['text'])
        except ValueError and TypeError as e: print(f"{repr(e)} [i={i} '{t}']")
        if ftext is None: continue
        if t["tag"]=="org": ftexts1.append(ftext)
        elif t["tag"]=="now": ftexts2.append(ftext)
        else: raise ValueError(f"unknow tag {t['tag']} [i={i} '{t}']")
    assert(len(ftexts1) == len(ftexts2))
    return save_ftext(ftexts1, ftexts2, outpath)

@filter_loadfiles(0)
def json2ftext(binobj: Union[str, List[str]], outpath=None) -> List[str]:
    ftexts1, ftexts2 = [], []
    for i, t in enumerate(json.loads(binobj)):
        if "org" in t:
            t2 = t["org"]
            ftext = ftext_t(int(t2['addr'], 0), int(t2['size'], 0), t2['text'])
            ftexts1.append(ftext)
        if "now" in t:
            t2 = t["now"]
            ftext = ftext_t(int(t2['addr'], 0), int(t2['size'], 0), t2['text'])
            ftexts2.append(ftext)
    assert(len(ftexts1) == len(ftexts2))
    return save_ftext(ftexts1, ftexts2, outpath)

@filter_loadfiles(0)
def paratranz2ftext(binobj: Union[str, List[str]], outpath=None) -> List[str]:
    ftexts1, ftexts2 = [], []
    for i, t in enumerate(json.loads(binobj)):
        text1, text2 = t["original"], t["translation"]
        k1, k2, k3 = t["key"].split("|")
        ftexts1.append(ftext_t(int(k2, 16), int(k3, 16), text1))
        ftexts2.append(ftext_t(int(k2, 16), int(k3, 16), text2 if text2!="" else text1))
    assert(len(ftexts1) == len(ftexts2))
    return save_ftext(ftexts1, ftexts2, outpath)

def docx2ftext(docxobj, outpath=None) -> List[str]:
    lines = []
    document = Document(docxobj)
    # text is the whole text, p.run are every str in styles
    for p in document.paragraphs:
        line = p.text.rstrip('\n').rstrip('\r') 
        lines.append(line + '\n')
    if outpath: writebytes(outpath, writelines(lines))
    return lines


def cli(cmdstr=None):
    def cmd_convert():
        flag = False
        if inpath_ext== '.txt': #  ftext to others
            if outpath_ext == '.txt': ftext2pretty(inpath, outpath)
            elif outpath_ext == '.docx': ftext2docx(inpath, outpath)
            elif outpath_ext == '.csv': ftext2csv(inpath, outpath)
            elif outpath_ext == ".json": ftext2json(inpath, outpath)
            elif outpath_ext == ".paratranz": ftext2paratranz(inpath, outpath)
            else: flag = True
        elif outpath_ext == '.txt': # others to ftext
            if inpath_ext == '.docx': docx2ftext(inpath, outpath)
            elif inpath_ext == '.csv': csv2ftext(inpath, outpath)
            elif inpath_ext == '.json': json2ftext(inpath, outpath)
            elif inpath_ext == ".paratranz": paratranz2ftext(inpath, outpath)
            else: flag = True
            return
        else: flag = True
        if not flag: return
        raise NotImplementedError(f"convert not support {inpath_ext}->{outpath_ext}")
    
    def cmd_split():
        outbase = os.path.splitext(args.outpath)[0]
        ftexts1, ftexts2 = load_ftext(args.inpath)
        nfile = args.split
        if nfile ==0: raise ValueError("nfile can not be 0")
        nftextall = len(ftexts1)
        nftexteach = (nftextall + nfile - 1) // nfile
        for i in range(nfile):
            outpath = outbase + f"_{i}.txt"
            s = slice(i*nftexteach, min((i+1)*nftexteach, nftextall))
            save_ftext(ftexts1[s], ftexts2[s], outpath)
    
    def cmd_merge():
        nfile = args.merge
        if nfile == 0:  inpaths = glob(args.inpath)
        else:  inpaths = [os.path.splitext(args.inpath)[0] + f"_{i}.txt" for i in range(nfile)]
        ftexts1, ftexts2 = [], []
        for inpath in inpaths:
            _f1, _f2 = load_ftext(inpath)
            ftexts1.extend(_f1); ftexts2.extend(_f2)
        save_ftext(ftexts1, ftexts2, args.outpath)

    parser = argparse.ArgumentParser(description=__description__)
    parser.add_argument("inpath")
    parser.add_argument("-o", "--outpath", default="out.txt")
    method = parser.add_mutually_exclusive_group()
    method.add_argument("--split", metavar="nfile", type=int, default=None)
    method.add_argument("--merge", metavar="nfile", type=int, default=None)
    method.add_argument("--convert", action="store_true")

    args = parser.parse_args(cmdstr.split(' ') if cmdstr else None)
    inpath, outpath = args.inpath, args.outpath
    outpath = outpath if len(outpath) > 0 else None
    inpath_ext = os.path.splitext(inpath)[1].lower()
    outpath_ext = os.path.splitext(outpath)[1].lower()
    if args.split is not None: cmd_split()
    elif args.merge is not None: cmd_merge()
    else: cmd_convert()

if __name__ == '__main__':
    cli()

"""
history:
v0.1, initial version with formatftext, docx2ftext, ftext2docx
v0.2, add support for csv and json, compatiable with paratranz.cn
v0.3, remake according to libtext v0.6
v0.3.1, add split merge ftext
v0.3.2, add paratranz json format
"""