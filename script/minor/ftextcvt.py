"""
This is a tool to change or adjust format 
in ftext made by bintext.py
    v0.1, developed by devseed
"""

import os
import codecs
import argparse
import docx # pip install python-docx
from docx.shared import Pt

def docx2ftext(wordpath, outpath=""):
    lines = []
    document = docx.Document(wordpath)
    for p in document.paragraphs:
        # text is the whole text, p.run are every str in styles
        line = p.text.rstrip('\n').rstrip('\r') 
        lines.append(line + '\n')
    if outpath!="": 
        with open(outpath, "wb") as fp:
            for line in lines:
                fp.write(line.encode('utf-8'))
    return lines

def ftext2docx(ftextpath, outpath=""):
    with codecs.open(ftextpath, 'r', 'utf-8') as fp:
        lines = fp.readlines()
    document = docx.Document()
    for i, line in enumerate(lines):
        line = line.rstrip('\n').rstrip('\r')
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)
    if outpath!="":
        document.save(outpath)
    return document

def formatftext(ftextpath, outpath=""):
    with codecs.open(ftextpath, 'r', 'utf-8') as fp:
        lines = fp.readlines()
    
    for i, line in enumerate(lines):
        _c = line[0]
        if _c=='○' or _c=='●':
            _idx = line.find(_c, 1) + 1
            if _idx==-1:
                raise ValueError(f"lines[{i}] error not closed, {line}!")
            elif line[_idx]!=' ': 
                print(f"detect line[{i}] without space")
                line = line[0:_idx] + ' ' + line[_idx: ]
        line = line.rstrip('\n').rstrip('\r')
        lines[i] = line + '\n'

    if outpath!="": 
        with open(outpath, "wb") as fp:
            for line in lines:
                fp.write(line.encode('utf-8'))
    return lines

def debug():
    pass

def main():
    parser = argparse.ArgumentParser(
        description="convert the ftext format to others, format is detected by outpath ext")
    parser.add_argument("inpath", type=str)
    parser.add_argument("-o", "--outpath", 
        type=str, default="out.txt")
    args = parser.parse_args()
    
    outpath = args.outpath
    outpath_ext=os.path.splitext(outpath)[1].lower()
    inpath = args.inpath
    inpath_ext = os.path.splitext(inpath)[1].lower()
    
    if inpath_ext== '.txt': #  ftext to others
        if outpath_ext == '.txt':
            formatftext(inpath, outpath)
        elif outpath_ext == '.docx':
            ftext2docx(inpath, outpath)
        else:
            raise NotImplementedError(f"unkonw format {inpath_ext}->{outpath_ext}")

    elif outpath_ext == '.txt': # others to ftext
        if inpath_ext == '.docx':
            docx2ftext(inpath, outpath)
        else:
            raise NotImplementedError(f"unkonw format {inpath_ext}->{outpath_ext}")
        return
    else:
        raise NotImplementedError(f"unkonw format {inpath_ext}->{outpath_ext}")

if __name__ == '__main__':
    #debug()
    main()
    pass

"""
history:
v0.1, initial version with formatftext, docx2ftext, ftext2docx
"""